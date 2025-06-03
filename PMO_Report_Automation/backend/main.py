from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse, FileResponse
from fastapi.staticfiles import StaticFiles
import io
import os
from typing import List
from urllib.parse import quote

# 導入所有核心邏輯和輔助函式
from .logic import (
    set_doc_normal_font,
    extract_text_list_from_ppt,
    extract_relevant_tables_from_ppt,
    process_work_summary_table,
    add_filtered_tables,
    insert_dynamic_meeting_section,
    add_legend_and_status
)
from docx import Document
from pptx import Presentation

# 初始化 FastAPI 應用程式
app = FastAPI(
    title="PMO 報告自動化 API",
    description="用於從 Word 範本和 PowerPoint 檔案自動生成 PMO 報告的 API。"
)

# ✅ 掛載 static 靜態檔案（提供 /static/favicon.png）
app.mount(
    "/static",
    StaticFiles(directory=os.path.join(os.path.dirname(__file__), "..", "static")),
    name="static"
)

# HTML 表單頁面路由
@app.get("/", response_class=HTMLResponse, summary="提供報告生成表單")
async def read_root():
    """
    提供用於上傳檔案和生成報告的主 HTML 頁面。
    """
    return FileResponse("frontend/index.html")


# 處理檔案上傳和報告生成
@app.post("/process-files/", summary="從上傳檔案生成 PMO 報告")
async def process_files(
    word_template: UploadFile = File(..., description="一個 Word (.docx) 範本檔案。"),
    ppt_files: List[UploadFile] = File(..., description="一個或多個 PowerPoint (.pptx) 檔案。")
):
    if not word_template.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Word 範本檔案必須是 .docx 格式。")

    for ppt_file in ppt_files:
        if not ppt_file.filename.lower().endswith(".pptx"):
            raise HTTPException(status_code=400, detail="PPT 檔案必須是 .pptx 格式。")

    try:
        word_template_stream = io.BytesIO(await word_template.read())
        doc = Document(word_template_stream)
        set_doc_normal_font(doc)

        for i, ppt_file in enumerate(ppt_files):
            print(f"\n--- 正在處理第 {i+1}/{len(ppt_files)} 個 PPT 文件：{ppt_file.filename} ---")
            add_page_break = (i > 0)

            ppt_stream = io.BytesIO(await ppt_file.read())
            ppt = Presentation(ppt_stream)

            print("💡 正在從 PPT 提取本月概要 (文字框內容)...")
            monthly_summary_items = extract_text_list_from_ppt(
                ppt,
                "本月概要",
                stop_keywords=["本期主要成果", "下期主要計畫", "專案階段", "執行狀況", "優化工作階段"]
            )

            print("💡 正在從 PPT 提取所有相關表格 (包含工作總覽, 專案階段, 程式修改)...")
            all_relevant_tables_data = extract_relevant_tables_from_ppt(ppt)

            work_summary_from_table = []
            if all_relevant_tables_data["work_summary_table_data"]:
                work_summary_from_table = process_work_summary_table(all_relevant_tables_data["work_summary_table_data"])
            else:
                print("ℹ️ 未找到工作總覽表格。")

            insert_dynamic_meeting_section(doc, ppt_file.filename, work_summary_from_table, monthly_summary_items, add_page_break)
            print("✅ 工作總覽與本月概要插入完成。")

            if all_relevant_tables_data["project_stage_tables_data"]:
                print("💡 正在插入專案階段表格...")
                add_filtered_tables(doc, all_relevant_tables_data["project_stage_tables_data"])
                doc.add_paragraph("")
                print("✅ 專案階段表格已插入 Word 文件。")
            else:
                print("ℹ️ 未找到專案階段表格。")

            if all_relevant_tables_data["program_modifications_table_data"]:
                print("💡 正在插入程式修改表格...")
                doc.add_heading("修改因重跑過程發現的問題所產生的程式修改", level=3)
                for table_data in all_relevant_tables_data["program_modifications_table_data"]:
                    add_filtered_tables(doc, [table_data])
                    doc.add_paragraph("")
                print("✅ 程式修改表格已插入 Word 文件。")
            else:
                print("ℹ️ 未找到符合條件的程式修改表格 (或該表格不在第二張投影片或缺少關鍵字)。")

            print("💡 正在添加圖例和狀態說明...")
            add_legend_and_status(doc)
            print("✅ 圖例和狀態說明已添加。")

        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        output_filename = "PMO_Consolidated_Report.docx"
        if ppt_files:
            first_ppt_name_base = os.path.splitext(ppt_files[0].filename)[0]
            output_filename = f"{first_ppt_name_base}_PMO_報告.docx"

        encoded_filename = quote(output_filename)

        return StreamingResponse(
            output_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
                "Content-Length": str(output_stream.getbuffer().nbytes)
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"處理檔案時發生錯誤：{e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"處理檔案時發生內部錯誤: {e}")
